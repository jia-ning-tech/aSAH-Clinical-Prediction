import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import joblib
import json
import os
from scipy import stats
from sklearn.metrics import roc_curve, auc, brier_score_loss, accuracy_score, confusion_matrix
from sklearn.calibration import calibration_curve

# 尝试导入 python-docx
try:
    from docx import Document
except ImportError:
    Document = None

plt.style.use('seaborn-v0_8-white')
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.titlesize'] = 14
plt.rcParams['axes.labelsize'] = 12
plt.rcParams['figure.dpi'] = 300

class PublicationReporter:
    def __init__(self, data_path, model_path):
        self.output_dir = "results"
        self.asset_dir = "app/assets" # 同时保存到前端目录
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.asset_dir, exist_ok=True)
        
        print(">> Loading resources...")
        self.raw_data = pd.read_csv("data/raw/clinical_data_raw.csv")
        self.processed_data = pd.read_csv(data_path)
        self.model_data = joblib.load(model_path)
        
        from sklearn.linear_model import LogisticRegression
        from sklearn.svm import SVC
        from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
        from sklearn.neural_network import MLPClassifier
        from xgboost import XGBClassifier
        
        # 定义模型 (参数需与 model_trainer 一致)
        self.models = {
            'LR': LogisticRegression(max_iter=5000, random_state=42),
            'SVM': SVC(probability=True, random_state=42),
            'GBM': GradientBoostingClassifier(random_state=42),
            'NN': MLPClassifier(hidden_layer_sizes=(100, 50), max_iter=1000, random_state=42),
            'RF': RandomForestClassifier(n_estimators=200, random_state=42),
            'XGBoost': XGBClassifier(eval_metric='logloss', random_state=42)
        }
        
        with open("data/meta/selected_features.json", "r") as f:
            self.features = json.load(f)['final_features']
            
        self.X_train = self.model_data['X_train_smote'][self.features]
        self.y_train = self.model_data['y_train_smote']
        self.X_test = self.model_data['X_test'][self.features]
        self.y_test = self.model_data['y_test']
        
        print(">> Retraining models for plotting (using R-selected features)...")
        for name, model in self.models.items():
            model.fit(self.X_train, self.y_train)

    def save_table_formats(self, df, filename_base):
        """通用保存函数"""
        base_path = os.path.join(self.output_dir, filename_base)
        df.to_csv(f"{base_path}.csv", index=False)
        try:
            with open(f"{base_path}.md", "w", encoding="utf-8") as f:
                f.write(df.to_markdown(index=False))
        except: pass

        if Document:
            try:
                doc = Document()
                doc.add_heading(filename_base, 0)
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                # 表头
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns): hdr_cells[i].text = str(col)
                # 内容
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row): row_cells[i].text = str(val)
                doc.save(f"{base_path}.docx")
            except: pass
        print(f"   Saved Table: {filename_base} (.csv, .md, .docx)")

    def generate_table_1(self):
        """Table 1: 基线特征表"""
        print("\n>> Generating Table 1...")
        df = self.raw_data.copy()
        
        # 自动识别列 (简单规则)
        feat_num = ['Age', 'GCS', 'Wfns', 'Hunthess', 'Modified fisher', 'PNI', 'Albumin', 'SIRI', 'SII', 'RBC', 'Hb', 'Ddimer', 'Fibrinogen', 'Sugar', 'AISI', 'PAR', 'NAR', 'PLR', 'MLR', 'CLR', 'NLPR', 'Procalcitonin']
        feat_cat = ['Sex', 'Hypertension', 'Smoke', 'Drink', 'Aneurysm location', 'Surgical method']
        
        feat_num = [f for f in feat_num if f in df.columns]
        feat_cat = [f for f in feat_cat if f in df.columns]
        
        group0 = df[df['Outcome'] == 0]
        group1 = df[df['Outcome'] == 1]
        n_total = len(df); n_g0 = len(group0); n_g1 = len(group1)
        
        rows = []
        # 连续变量
        for var in feat_num:
            row = {'Variable': f"{var} (mean ± SD)"}
            m_all = df[var].mean(); s_all = df[var].std()
            m0 = group0[var].mean(); s0 = group0[var].std()
            m1 = group1[var].mean(); s1 = group1[var].std()
            
            row['Overall'] = f"{m_all:.2f} ± {s_all:.2f}"
            row['Good Outcome'] = f"{m0:.2f} ± {s0:.2f}"
            row['Poor Outcome'] = f"{m1:.2f} ± {s1:.2f}"
            
            try:
                t, p = stats.ttest_ind(group0[var].dropna(), group1[var].dropna())
                row['P value'] = f"{p:.3f}" if p >= 0.001 else "<0.001"
            except: row['P value'] = "-"
            rows.append(row)
            
        # 分类变量
        for var in feat_cat:
            cats = sorted(df[var].dropna().unique())
            for cv in cats:
                row = {'Variable': f"{var}: {cv}"}
                cnt_all = len(df[df[var] == cv])
                cnt0 = len(group0[group0[var] == cv])
                cnt1 = len(group1[group1[var] == cv])
                
                row['Overall'] = f"{cnt_all} ({cnt_all/n_total*100:.1f}%)"
                row['Good Outcome'] = f"{cnt0} ({cnt0/n_g0*100:.1f}%)"
                row['Poor Outcome'] = f"{cnt1} ({cnt1/n_g1*100:.1f}%)"
                
                try:
                    c0_yes = cnt0; c0_no = n_g0 - cnt0
                    c1_yes = cnt1; c1_no = n_g1 - cnt1
                    chi2, p, _, _ = stats.chi2_contingency([[c0_yes, c0_no], [c1_yes, c1_no]])
                    row['P value'] = f"{p:.3f}" if p >= 0.001 else "<0.001"
                except: row['P value'] = "-"
                rows.append(row)

        self.save_table_formats(pd.DataFrame(rows), "Table1_Baseline")

    def generate_fig_4(self):
        """Fig 4: Feature Importance (6 Models)"""
        print("\n>> Generating Fig 4 (Feature Importance)...")
        fig, axes = plt.subplots(2, 3, figsize=(18, 10))
        axes = axes.flatten()
        model_names = ['LR', 'SVM', 'GBM', 'NN', 'RF', 'XGBoost']
        
        for i, name in enumerate(model_names):
            model = self.models[name]
            ax = axes[i]
            
            # 提取重要性
            importances = None
            if hasattr(model, 'feature_importances_'):
                importances = model.feature_importances_
            elif hasattr(model, 'coef_'):
                importances = np.abs(model.coef_[0])
            else:
                from sklearn.inspection import permutation_importance
                r = permutation_importance(model, self.X_test, self.y_test, n_repeats=5, random_state=42)
                importances = r.importances_mean

            # 排序 Top 10
            indices = np.argsort(importances)[::-1][:10]
            top_feats = [self.features[j] for j in indices]
            top_imps = importances[indices]
            
            # 归一化
            if top_imps.max() > 0: top_imps = 100 * (top_imps / top_imps.max())
            
            y_pos = np.arange(len(top_feats))
            ax.barh(y_pos, top_imps, align='center', color='#4c72b0')
            ax.set_yticks(y_pos)
            ax.set_yticklabels(top_feats)
            ax.invert_yaxis()
            ax.set_title(f"{chr(65+i)}. {name}")
            ax.set_xlabel('Relative Importance')
            
        plt.tight_layout()
        plt.savefig(f"{self.output_dir}/Fig4_FeatureImportance.png")
        print(f"   Saved: results/Fig4_FeatureImportance.png")

    def generate_fig_5_and_table_2(self):
        """Fig 5 (ROC/Calibration/DCA Panel) & Table 2"""
        print("\n>> Generating Fig 5 and Table 2...")
        
        # 这里的 2行 x 3列 (Train vs Test)
        fig, axes = plt.subplots(2, 3, figsize=(18, 12))
        datasets = [('Training Set', self.X_train, self.y_train), ('Test Set', self.X_test, self.y_test)]
        table2_rows = []
        
        for row_idx, (ds_name, X, y) in enumerate(datasets):
            ax_roc = axes[row_idx, 0]
            ax_cal = axes[row_idx, 1]
            ax_dca = axes[row_idx, 2]
            
            # 画基线
            ax_roc.plot([0, 1], [0, 1], 'k--', alpha=0.5)
            ax_cal.plot([0, 1], [0, 1], 'k:', label="Perfect")
            
            # DCA 基线
            thresholds = np.linspace(0.01, 0.99, 100)
            prevalence = y.mean()
            net_benefit_all = prevalence - (1-prevalence)*thresholds/(1-thresholds)
            ax_dca.plot(thresholds, net_benefit_all, 'k--', label='Treat All', alpha=0.5)
            ax_dca.plot(thresholds, np.zeros_like(thresholds), 'k-', label='Treat None', alpha=0.5)
            ax_dca.set_ylim(-0.05, 0.4)
            
            for name, model in self.models.items():
                probs = model.predict_proba(X)[:, 1]
                preds = model.predict(X)
                
                # Metrics
                fpr, tpr, _ = roc_curve(y, probs)
                roc_auc = auc(fpr, tpr)
                acc = accuracy_score(y, preds)
                tn, fp, fn, tp = confusion_matrix(y, preds).ravel()
                sens = tp / (tp + fn)
                spec = tn / (tn + fp)
                brier = brier_score_loss(y, probs)
                
                table2_rows.append({
                    'Dataset': ds_name, 'Model': name, 
                    'AUC': f"{roc_auc:.3f}", 'Accuracy': f"{acc:.3f}", 
                    'Sensitivity': f"{sens:.3f}", 'Specificity': f"{spec:.3f}", 
                    'Brier Score': f"{brier:.3f}"
                })
                
                # Plot ROC
                ax_roc.plot(fpr, tpr, label=f"{name} (AUC={roc_auc:.2f})")
                
                # Plot Calibration
                prob_true, prob_pred = calibration_curve(y, probs, n_bins=10)
                ax_cal.plot(prob_pred, prob_true, marker='.', label=name)
                
                # Plot DCA
                nb = []
                for thresh in thresholds:
                    tp_ = np.sum((probs >= thresh) & (y == 1))
                    fp_ = np.sum((probs >= thresh) & (y == 0))
                    n = len(y)
                    val = (tp_/n) - (fp_/n)*(thresh/(1-thresh))
                    nb.append(val)
                ax_dca.plot(thresholds, nb, label=name)
                
            ax_roc.set_title(f"{ds_name} - ROC"); ax_roc.legend(fontsize=8)
            ax_roc.set_xlabel('1 - Specificity'); ax_roc.set_ylabel('Sensitivity')
            
            ax_cal.set_title(f"{ds_name} - Calibration")
            ax_cal.set_xlabel('Predicted Probability'); ax_cal.set_ylabel('Observed Fraction')
            
            ax_dca.set_title(f"{ds_name} - DCA")
            ax_dca.set_xlabel('Threshold Probability'); ax_dca.set_ylabel('Net Benefit')
            
        plt.tight_layout()
        save_name = "Fig5_Performance_Panel.png"
        plt.savefig(f"{self.output_dir}/{save_name}")
        # 同时保存到 app 目录，修复前端不显示问题
        plt.savefig(f"{self.asset_dir}/roc_curves.png") # 暂时代替
        plt.savefig(f"{self.asset_dir}/calibration_curve.png")
        plt.savefig(f"{self.asset_dir}/dca_curve.png")
        print(f"   Saved: results/{save_name}")
        
        self.save_table_formats(pd.DataFrame(table2_rows), "Table2_Performance")

    def generate_fig_6(self):
        """Fig 6: SHAP Summary (XGBoost)"""
        print("\n>> Generating Fig 6 (SHAP)...")
        import shap
        model = self.models['XGBoost']
        explainer = shap.TreeExplainer(model)
        shap_values = explainer.shap_values(self.X_train)
        
        plt.figure(figsize=(10, 8))
        shap.summary_plot(shap_values, self.X_train, show=False, plot_type="dot")
        plt.tight_layout()
        plt.savefig(f"{self.output_dir}/Fig6_SHAP_Beeswarm.png")
        plt.savefig(f"{self.asset_dir}/shap_summary.png")
        print(f"   Saved: results/Fig6_SHAP_Beeswarm.png")

    def run_all(self):
        # ✅ 这次全部解锁！
        self.generate_table_1()
        self.generate_fig_4()
        self.generate_fig_5_and_table_2()
        self.generate_fig_6()
        print("\n✅ 所有图表复现完成！请查看 results/ 目录。")

if __name__ == "__main__":
    reporter = PublicationReporter(
        data_path="data/processed/modeling_data.csv",
        model_path="data/processed/train_test_data.pkl"
    )
    reporter.run_all()
